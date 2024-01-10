import os
import docx
from PyPDF2 import PdfFileWriter, PdfFileReader

def convert_to_pdf(file):
    doc = docx.Document(file)
    filename, ext = os.path.splitext(file)
    for index, page in enumerate(doc.pages):
        page_doc = docx.Document()
        page_doc.add_paragraph(page.text)
        page_doc.save(f'{filename}_{index}.docx')
        pdf_file=f'{filename}_{index}.pdf'
        writer = PdfFileWriter()
        docx_file = f'{filename}_{index}.docx'
        reader = docx.Document(docx_file)
        writer.addPage(reader.getPage(0))
        with open(pdf_file, 'wb') as f:
              writer.write(f)
        os.remove(docx_file)

if __name__=='__main__':
    word_file = 'sample.docx'
    convert_to_pdf(word_file)